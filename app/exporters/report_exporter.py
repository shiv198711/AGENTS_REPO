"""Report exporter — Markdown / JSON / DOCX (best-effort) outputs."""
from __future__ import annotations

import io
import json
from typing import Any

from ..models.execution import ExecutionRecord


class ReportExporter:
    def markdown(self, record: ExecutionRecord) -> str:
        lines: list[str] = []
        lines.append(f"# CVI/MDS SAP Note Implementation Report")
        lines.append("")
        lines.append(f"- Execution ID: `{record.execution_id}`")
        lines.append(f"- System: **{record.system_type.value}** / {record.system_tier.value}")
        lines.append(f"- Status: **{record.status.value}**")
        lines.append(f"- Requested by: {record.requested_by}")
        lines.append(f"- Started at: {record.started_at}")
        lines.append(f"- Finished at: {record.finished_at or '(pending)'}")
        lines.append("")

        # Notes
        lines.append("## SAP Notes")
        for n, m, v in zip(record.notes, record.metadata, record.validation):
            lines.append(
                f"- **{n.note_number}** — {m.title or n.title or '(no title)'} "
                f"| component={m.component or 'n/a'} | valid={v.valid} | implement={v.can_implement}"
            )
            if v.errors:
                for e in v.errors:
                    lines.append(f"    - ❌ {e}")
            if v.warnings:
                for w in v.warnings:
                    lines.append(f"    - ⚠️ {w}")
        lines.append("")

        # Analysis
        if record.analysis:
            lines.append("## AI Analysis")
            lines.append("")
            lines.append("### Executive Summary")
            lines.append(record.analysis.executive_summary or "(none)")
            lines.append("")
            lines.append("### Business Impact")
            lines.append(record.analysis.business_impact or "(none)")
            lines.append("")
            lines.append("### Technical Impact")
            lines.append(record.analysis.technical_impact or "(none)")
            lines.append("")
            if record.analysis.implementation_sequence:
                lines.append("### Implementation Sequence")
                for s in record.analysis.implementation_sequence:
                    lines.append(f"- {s}")
                lines.append("")
            if record.analysis.predicted_conflicts:
                lines.append("### Predicted Conflicts")
                for c in record.analysis.predicted_conflicts:
                    lines.append(f"- {c}")
                lines.append("")
            if record.analysis.recommended_prereqs:
                lines.append("### Recommended Prerequisite Notes")
                for p in record.analysis.recommended_prereqs:
                    lines.append(f"- {p}")
                lines.append("")

        # Transports
        if record.transports:
            lines.append("## Transports")
            for t in record.transports:
                lines.append(f"- `{t.request_id}` — {t.description} ({t.status})")
            lines.append("")

        # Cockpit
        if record.cockpit_result:
            cr = record.cockpit_result
            lines.append(f"## Post-Implementation Check ({cr.cockpit})")
            lines.append(f"- Passed: **{cr.passed}**")
            for k, v in cr.checks.items():
                lines.append(f"    - {k}: {v}")
            if cr.findings:
                lines.append("- Findings:")
                for f in cr.findings:
                    lines.append(f"    - {f}")
            if cr.remediation:
                lines.append("- Remediation:")
                for r in cr.remediation:
                    lines.append(f"    - {r}")
            lines.append("")

        # Logs
        if record.logs:
            lines.append("## Execution Log")
            for l in record.logs:
                lines.append(f"- [{l.ts}] [{l.level}] {l.step}: {l.message}")

        # Errors
        if record.error_summary:
            lines.append("")
            lines.append("## Errors")
            lines.append(record.error_summary)
        if record.rollback_recommendation:
            lines.append("")
            lines.append("## Rollback Recommendation")
            lines.append(record.rollback_recommendation)

        return "\n".join(lines).rstrip() + "\n"

    def json(self, record: ExecutionRecord) -> str:
        return json.dumps(record.model_dump(mode="json"), indent=2, default=str)

    def docx(self, record: ExecutionRecord) -> bytes:
        try:
            from docx import Document  # type: ignore

            doc = Document()
            doc.add_heading("CVI/MDS SAP Note Implementation Report", level=1)
            doc.add_paragraph(f"Execution ID: {record.execution_id}")
            doc.add_paragraph(
                f"System: {record.system_type.value} / {record.system_tier.value}"
            )
            doc.add_paragraph(f"Status: {record.status.value}")

            doc.add_heading("SAP Notes", level=2)
            for n, m in zip(record.notes, record.metadata):
                doc.add_paragraph(
                    f"{n.note_number} — {m.title or '(no title)'}",
                    style="List Bullet",
                )

            if record.analysis:
                doc.add_heading("AI Analysis", level=2)
                doc.add_paragraph(record.analysis.executive_summary or "")

            if record.cockpit_result:
                doc.add_heading(record.cockpit_result.cockpit, level=2)
                for k, v in record.cockpit_result.checks.items():
                    doc.add_paragraph(f"{k}: {v}", style="List Bullet")

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:  # noqa: BLE001
            # Fallback: return markdown as bytes if python-docx unavailable.
            return self.markdown(record).encode("utf-8")

    def summary(self, record: ExecutionRecord) -> dict[str, Any]:
        return {
            "execution_id": record.execution_id,
            "status": record.status,
            "system": record.system_type.value,
            "tier": record.system_tier.value,
            "notes": [n.note_number for n in record.notes],
            "transports": [
                {"trkorr": t.request_id, "status": t.status} for t in record.transports
            ],
            "cockpit_passed": bool(record.cockpit_result and record.cockpit_result.passed),
            "started_at": record.started_at,
            "finished_at": record.finished_at,
            "error_summary": record.error_summary,
        }